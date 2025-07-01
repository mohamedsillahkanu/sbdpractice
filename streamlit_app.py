import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import geopandas as gpd
import math
from io import BytesIO
import re

# Custom CSS for the dashboard
st.markdown("""
<style>
    .main .block-container {
        padding-top: 1rem;
        padding-bottom: 1rem;
        max-width: none;
    }
    
    h1 {
        color: #2c3e50;
        text-align: center;
        font-weight: 700;
        margin-bottom: 2rem;
    }
    
    h2, h3 {
        color: #34495e;
        border-bottom: 2px solid #3498db;
        padding-bottom: 0.5rem;
    }
    
    .stButton > button {
        background: linear-gradient(45deg, #3498db, #2980b9);
        color: white;
        border: none;
        border-radius: 25px;
        padding: 0.5rem 2rem;
        font-weight: 600;
    }
</style>
""", unsafe_allow_html=True)

def parse_gps_coordinates(gps_str):
    """Enhanced GPS coordinate parsing that handles multiple formats"""
    if pd.isna(gps_str):
        return None, None
    
    gps_str = str(gps_str).strip()
    lat, lon = None, None
    
    # Format 1: "8.6103181,-12.2029534"
    if ',' in gps_str and not ' ' in gps_str:
        try:
            parts = gps_str.split(',')
            if len(parts) == 2:
                lat = float(parts[0].strip())
                lon = float(parts[1].strip())
        except ValueError:
            pass
    
    # Format 2: "8.6103181 -12.2029534" (space separated)
    elif ' ' in gps_str and ',' not in gps_str:
        try:
            parts = gps_str.split()
            if len(parts) == 2:
                lat = float(parts[0].strip())
                lon = float(parts[1].strip())
        except ValueError:
            pass
    
    # Format 3: Other formats with parentheses, etc.
    else:
        # Extract numbers using regex
        numbers = re.findall(r'-?\d+\.?\d*', gps_str)
        if len(numbers) >= 2:
            try:
                lat = float(numbers[0])
                lon = float(numbers[1])
            except ValueError:
                pass
    
    return lat, lon

def create_chiefdom_mapping():
    """Create mapping between GPS data chiefdom names and shapefile FIRST_CHIE names"""
    chiefdom_mapping = {
        # BO District mappings
        "Bo City": "BO TOWN",
        "Badjia": "BADJIA",
        "Bargbo": "BAGBO",
        "Bagbwe": "BAGBWE(BAGBE)",
        "Baoma": "BOAMA",
        "Bongor": "BONGOR",
        "Bumpe Ngao": "BUMPE NGAO",
        "Gbo": "GBO",
        "Jaiama": "JAIAMA",
        "Kakua": "KAKUA",
        "Komboya": "KOMBOYA",
        "Lugbu": "LUGBU",
        "Niawa Lenga": "NIAWA LENGA",
        "Selenga": "SELENGA",
        "Tikonko": "TIKONKO",
        "Valunia": "VALUNIA",
        "Wonde": "WONDE",
        
        # BOMBALI District mappings
        "Biriwa": "BIRIWA",
        "Bombali Sebora": "BOMBALI SEBORA",
        "Bombali Serry": "BOMBALI SIARI",
        "Gbanti (Bombali)": "GBANTI",
        "Gbanti": "GBANTI",
        "Gbendembu": "GBENDEMBU",
        "Kamaranka": "KAMARANKA",
        "Magbaimba Ndohahun": "MAGBAIMBA NDORWAHUN",
        "Makarie": "MAKARI",
        "Mara": "MARA",
        "Ngowahun": "N'GOWAHUN",
        "Paki Masabong": "PAKI MASABONG",
        "Safroko Limba": "SAFROKO LIMBA",
        "Makeni City": "MAKENI CITY",
        
        # Add more mappings as needed
    }
    return chiefdom_mapping

def map_chiefdom_name(chiefdom_name, mapping):
    """Map chiefdom name from GPS data to shapefile name"""
    if pd.isna(chiefdom_name):
        return None
    
    chiefdom_name = str(chiefdom_name).strip()
    
    # Direct match
    if chiefdom_name in mapping:
        return mapping[chiefdom_name]
    
    # Case-insensitive match
    for key, value in mapping.items():
        if key.upper() == chiefdom_name.upper():
            return value
    
    # Partial match (contains)
    for key, value in mapping.items():
        if key.upper() in chiefdom_name.upper() or chiefdom_name.upper() in key.upper():
            return value
    
    # Return original if no mapping found
    return chiefdom_name
    """Enhanced GPS coordinate parsing that handles multiple formats"""
    if pd.isna(gps_str):
        return None, None
    
    gps_str = str(gps_str).strip()
    lat, lon = None, None
    
    # Format 1: "8.6103181,-12.2029534"
    if ',' in gps_str and not ' ' in gps_str:
        try:
            parts = gps_str.split(',')
            if len(parts) == 2:
                lat = float(parts[0].strip())
                lon = float(parts[1].strip())
        except ValueError:
            pass
    
    # Format 2: "8.6103181 -12.2029534" (space separated)
    elif ' ' in gps_str and ',' not in gps_str:
        try:
            parts = gps_str.split()
            if len(parts) == 2:
                lat = float(parts[0].strip())
                lon = float(parts[1].strip())
        except ValueError:
            pass
    
    # Format 3: Other formats with parentheses, etc.
    else:
        # Extract numbers using regex
        numbers = re.findall(r'-?\d+\.?\d*', gps_str)
        if len(numbers) >= 2:
            try:
                lat = float(numbers[0])
                lon = float(numbers[1])
            except ValueError:
                pass
    
    return lat, lon

def extract_gps_data_from_excel(df):
    """Extract GPS data from the Excel file"""
    # Create empty lists to store extracted data
    districts, chiefdoms, gps_locations = [], [], []
    
    # Get chiefdom mapping
    chiefdom_mapping = create_chiefdom_mapping()
    
    # Process each row in the "Scan QR code" column
    for idx, qr_text in enumerate(df["Scan QR code"]):
        if pd.isna(qr_text):
            districts.append(None)
            chiefdoms.append(None)
            gps_locations.append(None)
            continue
            
        # Extract values using regex patterns
        district_match = re.search(r"District:\s*([^\n]+)", str(qr_text))
        district = district_match.group(1).strip() if district_match else None
        districts.append(district)
        
        chiefdom_match = re.search(r"Chiefdom:\s*([^\n]+)", str(qr_text))
        original_chiefdom = chiefdom_match.group(1).strip() if chiefdom_match else None
        
        # Map chiefdom name to match shapefile
        mapped_chiefdom = map_chiefdom_name(original_chiefdom, chiefdom_mapping)
        chiefdoms.append(mapped_chiefdom)
        
        # Get GPS Location from the corresponding row
        if "GPS Location" in df.columns:
            gps_locations.append(df["GPS Location"].iloc[idx])
        else:
            gps_locations.append(None)
    
    # Create a new DataFrame with extracted values
    extracted_df = pd.DataFrame({
        "District": districts,
        "Chiefdom": chiefdoms,
        "GPS_Location": gps_locations
    })
    
    return extracted_df

def create_chiefdom_subplot_dashboard(gdf, extracted_df, district_name, cols=4):
    """Create subplot dashboard for all chiefdoms in a district"""
    
    # Filter shapefile for the district
    district_gdf = gdf[gdf['FIRST_DNAM'] == district_name].copy()
    
    if len(district_gdf) == 0:
        st.error(f"No chiefdoms found for {district_name} district in shapefile")
        return None
    
    # Get unique chiefdoms from shapefile
    chiefdoms = sorted(district_gdf['FIRST_CHIE'].dropna().unique())
    
    # Calculate rows needed
    rows = math.ceil(len(chiefdoms) / cols)
    
    # Create subplot figure with increased vertical space
    fig, axes = plt.subplots(rows, cols, figsize=(cols*5, rows*6))
    fig.suptitle(f'{district_name} District - All Chiefdoms with GPS Locations', 
                 fontsize=20, fontweight='bold', y=0.98)
    
    # Ensure axes is always 2D array
    if rows == 1:
        axes = axes.reshape(1, -1)
    elif cols == 1:
        axes = axes.reshape(-1, 1)
    
    # Plot each chiefdom
    for idx, chiefdom in enumerate(chiefdoms):
        row = idx // cols
        col = idx % cols
        ax = axes[row, col]
        
        # Filter shapefile for this specific chiefdom
        chiefdom_gdf = district_gdf[district_gdf['FIRST_CHIE'] == chiefdom].copy()
        
        # Plot chiefdom boundary
        chiefdom_gdf.plot(ax=ax, color='lightblue', edgecolor='navy', alpha=0.7, linewidth=2)
        
        # Filter GPS data for this district and chiefdom with exact matching
        district_data = extracted_df[extracted_df["District"].str.upper() == district_name.upper()].copy()
        chiefdom_data = district_data[district_data["Chiefdom"] == chiefdom].copy()
        
        # Extract and plot GPS coordinates for this chiefdom
        coords_extracted = []
        if len(chiefdom_data) > 0 and not chiefdom_data["GPS_Location"].isna().all():
            gps_data = chiefdom_data["GPS_Location"].dropna()
            
            for gps_val in gps_data:
                if pd.notna(gps_val):
                    lat, lon = parse_gps_coordinates(gps_val)
                    
                    # Validate coordinates for Sierra Leone
                    if lat is not None and lon is not None:
                        if 6.0 <= lat <= 11.0 and -14.0 <= lon <= -10.0:
                            coords_extracted.append([lat, lon, str(gps_val)])  # Include original GPS string for debugging
        
        # Handle overlapping coordinates by adding small offsets
        def separate_overlapping_points(coords, min_distance=0.001):
            """Separate overlapping GPS points by adding small offsets"""
            if len(coords) <= 1:
                return coords
            
            separated_coords = []
            for i, (lat, lon, gps_str) in enumerate(coords):
                adjusted_lat, adjusted_lon = lat, lon
                
                # Check if this point overlaps with any previous point
                for j, (prev_lat, prev_lon, _) in enumerate(separated_coords):
                    distance = ((adjusted_lat - prev_lat)**2 + (adjusted_lon - prev_lon)**2)**0.5
                    if distance < min_distance:
                        # Add small offset in a circular pattern
                        angle = (i * 2 * 3.14159) / len(coords)
                        offset = min_distance * 1.5
                        adjusted_lat += offset * math.cos(angle)
                        adjusted_lon += offset * math.sin(angle)
                
                separated_coords.append([adjusted_lat, adjusted_lon, gps_str])
            
            return separated_coords
        
        # Separate overlapping points
        if coords_extracted:
            coords_extracted = separate_overlapping_points(coords_extracted)
        
        # Plot GPS points if available
        if coords_extracted:
            lats, lons, gps_strings = zip(*coords_extracted)
            
            # Plot each point individually to ensure they're all visible
            for i, (lat, lon) in enumerate(zip(lats, lons)):
                ax.scatter(lon, lat, c='red', s=100, alpha=1.0, 
                          edgecolors='white', linewidth=2, zorder=100, marker='o')
            
            # Add debug info in title if multiple points exist
            if len(coords_extracted) > 1:
                unique_coords = len(set([(round(lat, 6), round(lon, 6)) for lat, lon, _ in coords_extracted]))
                if unique_coords < len(coords_extracted):
                    debug_info = f" [Debug: {len(coords_extracted)} total, {unique_coords} unique locations]"
                else:
                    debug_info = ""
            else:
                debug_info = ""
        
        # Set title and clean up axes
        ax.set_title(f'{chiefdom}\n({len(coords_extracted)} schools{debug_info if coords_extracted else ""})', 
                    fontsize=12, fontweight='bold', pad=10)
        
        # Remove axis labels and ticks for cleaner look
        ax.set_xticks([])
        ax.set_yticks([])
        ax.set_xlabel('')
        ax.set_ylabel('')
        
        # Remove the box frame
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.spines['bottom'].set_visible(False)
        ax.spines['left'].set_visible(False)
        
        # Add grid
        ax.grid(True, alpha=0.3, linestyle='--')
        
        # Set equal aspect ratio and tight layout
        ax.set_aspect('equal')
        
        # Set bounds to chiefdom extent with some padding
        bounds = chiefdom_gdf.total_bounds
        padding = 0.01
        ax.set_xlim(bounds[0] - padding, bounds[2] + padding)
        ax.set_ylim(bounds[1] - padding, bounds[3] + padding)
    
    # Hide empty subplots
    total_plots = rows * cols
    for idx in range(len(chiefdoms), total_plots):
        row = idx // cols
        col = idx % cols
        axes[row, col].set_visible(False)
    
    plt.tight_layout()
    plt.subplots_adjust(top=0.93, hspace=0.4, wspace=0.3)
    
    return fig

# Streamlit App
st.title("🗺️ Chiefdom GPS Dashboard - BO and BOMBALI Districts")
st.markdown("**Comprehensive view of all chiefdoms with GPS school locations**")

# Load the data
try:
    df_original = pd.read_excel("GMB253374_SBD_ITN_clean.xlsx")
    
    # Extract GPS data with chiefdom mapping
    extracted_df = extract_gps_data_from_excel(df_original)
    
except Exception as e:
    st.error(f"❌ Error loading Excel file: {e}")
    st.stop()

# Load shapefile
try:
    gdf = gpd.read_file("Chiefdom2021.shp")
    
except Exception as e:
    st.error(f"❌ Could not load shapefile: {e}")
    st.stop()

# Dashboard Settings
columns = 4
show_data_info = False

if show_data_info:
    # Display data information
    st.subheader("📊 Data Overview")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        total_records = len(extracted_df)
        st.metric("Total Records", f"{total_records:,}")
    
    with col2:
        bo_records = len(extracted_df[extracted_df["District"].str.upper() == "BO"])
        st.metric("BO District", f"{bo_records:,}")
    
    with col3:
        bombali_records = len(extracted_df[extracted_df["District"].str.upper() == "BOMBALI"])
        st.metric("BOMBALI District", f"{bombali_records:,}")
    
    with col4:
        total_gps = len(extracted_df[extracted_df["GPS_Location"].notna()])
        st.metric("GPS Records", f"{total_gps:,}")

# Create dashboards
st.header("🗺️ District Dashboards")

# BO District Dashboard
st.subheader("1a. BO District - All Chiefdoms")

with st.spinner("Generating BO District dashboard..."):
    try:
        fig_bo = create_chiefdom_subplot_dashboard(gdf, extracted_df, "BO", columns)
        if fig_bo:
            st.pyplot(fig_bo)
            
            # Save figure option
            buffer_bo = BytesIO()
            fig_bo.savefig(buffer_bo, format='png', dpi=300, bbox_inches='tight')
            buffer_bo.seek(0)
            
            st.download_button(
                label="📥 Download BO District Dashboard",
                data=buffer_bo,
                file_name="BO_District_Chiefdoms_Dashboard.png",
                mime="image/png"
            )
        else:
            st.warning("Could not generate BO District dashboard")
    except Exception as e:
        st.error(f"Error generating BO District dashboard: {e}")

st.divider()

# BOMBALI District Dashboard
st.subheader("1b. BOMBALI District - All Chiefdoms")

with st.spinner("Generating BOMBALI District dashboard..."):
    try:
        fig_bombali = create_chiefdom_subplot_dashboard(gdf, extracted_df, "BOMBALI", columns)
        if fig_bombali:
            st.pyplot(fig_bombali)
            
            # Save figure option
            buffer_bombali = BytesIO()
            fig_bombali.savefig(buffer_bombali, format='png', dpi=300, bbox_inches='tight')
            buffer_bombali.seek(0)
            
            st.download_button(
                label="📥 Download BOMBALI District Dashboard",
                data=buffer_bombali,
                file_name="BOMBALI_District_Chiefdoms_Dashboard.png",
                mime="image/png"
            )
        else:
            st.warning("Could not generate BOMBALI District dashboard")
    except Exception as e:
        st.error(f"Error generating BOMBALI District dashboard: {e}")

# Summary Statistics
st.header("📈 Summary Statistics")

# Create summary for both districts
summary_data = []

for district in ["BO", "BOMBALI"]:
    district_gdf = gdf[gdf['FIRST_DNAM'] == district]
    district_data = extracted_df[extracted_df["District"].str.upper() == district.upper()]
    
    # Count GPS locations
    gps_count = len(district_data[district_data["GPS_Location"].notna()])
    
    summary_data.append({
        "District": district,
        "Chiefdoms": len(district_gdf) if len(district_gdf) > 0 else 0,
        "Total Records": len(district_data),
        "GPS Records": gps_count,
        "GPS Coverage": f"{(gps_count/len(district_data)*100) if len(district_data) > 0 else 0:.1f}%"
    })

summary_df = pd.DataFrame(summary_data)
st.dataframe(summary_df, use_container_width=True)

# Section 2: School Coverage Dashboard
st.header("📊 Section 2: School Coverage Analysis")

def generate_target_school_data(chiefdoms):
    """Generate target school data based on real chiefdom data"""
    
    # Real target data provided
    target_data = {
        # BO District
        "BADJIA": 9,
        "BAGBWE(BAGBE)": 18,
        "BOAMA": 56,
        "BAGBO": 31,  # Bargbo
        "BO TOWN": 86,  # Bo City
        "BONGOR": 18,
        "BUMPE NGAO": 63,  # Bumpeh
        "GBO": 10,
        "JAIAMA": 25,
        "KAKUA": 164,
        "KOMBOYA": 17,
        "LUGBU": 32,
        "NIAWA LENGA": 25,
        "SELENGA": 7,
        "TIKONKO": 89,  # Tinkoko
        "VALUNIA": 38,
        "WONDE": 13,
        
        # BOMBALI District  
        "BIRIWA": 48,
        "BOMBALI SEBORA": 44,
        "BOMBALI SIARI": 7,  # Bombali Serry Chiefdom
        "GBANTI": 40,  # Gbanti (Bombali)
        "GBENDEMBU": 30,
        "KAMARANKA": 13,
        "MAGBAIMBA NDORWAHUN": 17,  # Magbaimba Ndohahun
        "MAKARI": 54,  # Makarie
        "MAKENI CITY": 93,
        "MARA": 15,
        "N'GOWAHUN": 28,  # Ngowahun
        "PAKI MASABONG": 29,
        "SAFROKO LIMBA": 36,
    }
    
    # For any chiefdom not in the list, return a default value
    result = {}
    for chiefdom in chiefdoms:
        result[chiefdom] = target_data.get(chiefdom, 20)  # Default to 20 if not found
    
    return result

def get_coverage_color(coverage_percent):
    """Get color based on coverage percentage"""
    if coverage_percent < 20:
        return '#d32f2f'  # Red
    elif coverage_percent < 40:
        return '#f57c00'  # Orange
    elif coverage_percent < 60:
        return '#fbc02d'  # Yellow
    elif coverage_percent < 80:
        return '#388e3c'  # Light Green
    elif coverage_percent < 100:
        return '#1976d2'  # Blue
    else:
        return '#4a148c'  # Purple (100% coverage)

def create_coverage_dashboard(gdf, extracted_df, district_name, cols=4):
    """Create coverage dashboard for all chiefdoms in a district"""
    
    # Filter shapefile for the district
    district_gdf = gdf[gdf['FIRST_DNAM'] == district_name].copy()
    
    if len(district_gdf) == 0:
        st.error(f"No chiefdoms found for {district_name} district in shapefile")
        return None
    
    # Get unique chiefdoms from shapefile
    chiefdoms = sorted(district_gdf['FIRST_CHIE'].dropna().unique())
    
    # Generate real target data
    target_data = generate_target_school_data(chiefdoms)
    
    # Calculate rows needed
    rows = math.ceil(len(chiefdoms) / cols)
    
    # Create subplot figure with increased vertical space
    fig, axes = plt.subplots(rows, cols, figsize=(cols*5, rows*6))
    fig.suptitle(f'{district_name} District - School Coverage Analysis', 
                 fontsize=20, fontweight='bold', y=0.98)
    
    # Ensure axes is always 2D array
    if rows == 1:
        axes = axes.reshape(1, -1)
    elif cols == 1:
        axes = axes.reshape(-1, 1)
    
    # Plot each chiefdom
    for idx, chiefdom in enumerate(chiefdoms):
        row = idx // cols
        col = idx % cols
        ax = axes[row, col]
        
        # Filter shapefile for this specific chiefdom
        chiefdom_gdf = district_gdf[district_gdf['FIRST_CHIE'] == chiefdom].copy()
        
        # Filter GPS data for this district and chiefdom with exact matching
        district_data = extracted_df[extracted_df["District"].str.upper() == district_name.upper()].copy()
        chiefdom_data = district_data[district_data["Chiefdom"] == chiefdom].copy()
        
        # Count actual schools (records) for this chiefdom
        actual_schools = len(chiefdom_data)
        
        # Get target schools from dummy data
        target_schools = target_data.get(chiefdom, 50)  # Default to 50 if not found
        
        # Calculate coverage percentage
        coverage_percent = (actual_schools / target_schools * 100) if target_schools > 0 else 0
        coverage_percent = min(coverage_percent, 100)  # Cap at 100%
        
        # Get color based on coverage
        coverage_color = get_coverage_color(coverage_percent)
        
        # Plot chiefdom boundary with coverage color
        chiefdom_gdf.plot(ax=ax, color=coverage_color, edgecolor='black', alpha=0.8, linewidth=2)
        
        # Create coverage text
        coverage_text = f"{actual_schools}/{target_schools} ({coverage_percent:.0f}%)"
        
        # Set title with coverage information
        ax.set_title(f'{chiefdom}\n{coverage_text}', 
                    fontsize=11, fontweight='bold', pad=10)
        
        # Remove axis labels and ticks for cleaner look
        ax.set_xticks([])
        ax.set_yticks([])
        ax.set_xlabel('')
        ax.set_ylabel('')
        
        # Remove the box frame
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.spines['bottom'].set_visible(False)
        ax.spines['left'].set_visible(False)
        
        # Add light grid
        ax.grid(True, alpha=0.2, linestyle='--')
        
        # Set equal aspect ratio and tight layout
        ax.set_aspect('equal')
        
        # Set bounds to chiefdom extent with some padding
        bounds = chiefdom_gdf.total_bounds
        padding = 0.01
        ax.set_xlim(bounds[0] - padding, bounds[2] + padding)
        ax.set_ylim(bounds[1] - padding, bounds[3] + padding)
    
    # Hide empty subplots
    total_plots = rows * cols
    for idx in range(len(chiefdoms), total_plots):
        row = idx // cols
        col = idx % cols
        axes[row, col].set_visible(False)
    
    plt.tight_layout()
    plt.subplots_adjust(top=0.93, hspace=0.4, wspace=0.3)
    
    return fig

# BO District Coverage Dashboard
st.subheader("2a. BO District - School Coverage")

with st.spinner("Generating BO District coverage dashboard..."):
    try:
        fig_bo_coverage = create_coverage_dashboard(gdf, extracted_df, "BO", columns)
        if fig_bo_coverage:
            st.pyplot(fig_bo_coverage)
            
            # Save figure option
            buffer_bo_coverage = BytesIO()
            fig_bo_coverage.savefig(buffer_bo_coverage, format='png', dpi=300, bbox_inches='tight')
            buffer_bo_coverage.seek(0)
            
            st.download_button(
                label="📥 Download BO District Coverage Dashboard",
                data=buffer_bo_coverage,
                file_name="BO_District_Coverage_Dashboard.png",
                mime="image/png"
            )
        else:
            st.warning("Could not generate BO District coverage dashboard")
    except Exception as e:
        st.error(f"Error generating BO District coverage dashboard: {e}")

st.divider()

# BOMBALI District Coverage Dashboard
st.subheader("2b. BOMBALI District - School Coverage")

with st.spinner("Generating BOMBALI District coverage dashboard..."):
    try:
        fig_bombali_coverage = create_coverage_dashboard(gdf, extracted_df, "BOMBALI", columns)
        if fig_bombali_coverage:
            st.pyplot(fig_bombali_coverage)
            
            # Save figure option
            buffer_bombali_coverage = BytesIO()
            fig_bombali_coverage.savefig(buffer_bombali_coverage, format='png', dpi=300, bbox_inches='tight')
            buffer_bombali_coverage.seek(0)
            
            st.download_button(
                label="📥 Download BOMBALI District Coverage Dashboard",
                data=buffer_bombali_coverage,
                file_name="BOMBALI_District_Coverage_Dashboard.png",
                mime="image/png"
            )
        else:
            st.warning("Could not generate BOMBALI District coverage dashboard")
    except Exception as e:
        st.error(f"Error generating BOMBALI District coverage dashboard: {e}")

# Section 3: ITN Coverage Analysis
st.header("🛡️ Section 3: ITN Coverage Analysis")

def extract_itn_data_from_excel(df):
    """Extract ITN coverage data from the Excel file"""
    # Create empty lists to store extracted data
    districts, chiefdoms, total_enrollment, total_itns, distributed_itns = [], [], [], [], []
    
    # Get chiefdom mapping
    chiefdom_mapping = create_chiefdom_mapping()
    
    # Process each row in the "Scan QR code" column
    for idx, qr_text in enumerate(df["Scan QR code"]):
        if pd.isna(qr_text):
            districts.append(None)
            chiefdoms.append(None)
            total_enrollment.append(0)
            total_itns.append(0)
            continue
            
        # Extract values using regex patterns
        district_match = re.search(r"District:\s*([^\n]+)", str(qr_text))
        district = district_match.group(1).strip() if district_match else None
        districts.append(district)
        
        chiefdom_match = re.search(r"Chiefdom:\s*([^\n]+)", str(qr_text))
        original_chiefdom = chiefdom_match.group(1).strip() if chiefdom_match else None
        
        # Map chiefdom name to match shapefile
        mapped_chiefdom = map_chiefdom_name(original_chiefdom, chiefdom_mapping)
        chiefdoms.append(mapped_chiefdom)
        
        # Calculate total enrollment (sum of all class enrollments)
        enrollment_total = 0
        for class_num in range(1, 6):  # Classes 1-5
            enrollment_col = f"How many pupils are enrolled in Class {class_num}?"
            if enrollment_col in df.columns:
                class_enrollment = df[enrollment_col].iloc[idx]
                if pd.notna(class_enrollment):
                    enrollment_total += int(class_enrollment)
        
        total_enrollment.append(enrollment_total)
        
        # Calculate total ITNs distributed (boys + girls only)
        itns_distributed = 0
        for class_num in range(1, 6):  # Classes 1-5
            # ITNs distributed (boys + girls)
            boys_col = f"How many boys in Class {class_num} received ITNs?"
            girls_col = f"How many girls in Class {class_num} received ITNs?"
            
            if boys_col in df.columns:
                boys_itns = df[boys_col].iloc[idx]
                if pd.notna(boys_itns):
                    itns_distributed += int(boys_itns)
            
            if girls_col in df.columns:
                girls_itns = df[girls_col].iloc[idx]
                if pd.notna(girls_itns):
                    itns_distributed += int(girls_itns)
        
        # Calculate total ITNs (distributed + left at school)
        itns_total = itns_distributed
        for class_num in range(1, 6):  # Classes 1-5
            # ITNs left at school for absent pupils - using correct column name
            left_col = "ITNs left at the school for pupils who were absent."
            if left_col in df.columns:
                left_itns = df[left_col].iloc[idx]
                if pd.notna(left_itns):
                    itns_total += int(left_itns)
        
        total_itns.append(itns_total)
        distributed_itns.append(itns_distributed)
    
    # Create a new DataFrame with extracted values
    itn_df = pd.DataFrame({
        "District": districts,
        "Chiefdom": chiefdoms,
        "Total_Enrollment": total_enrollment,
        "Total_ITNs": total_itns,
        "Distributed_ITNs": distributed_itns
    })
    
    return itn_df

def create_itn_coverage_dashboard(gdf, itn_df, district_name, cols=4):
    """Create ITN coverage dashboard for all chiefdoms in a district"""
    
    # Filter shapefile for the district
    district_gdf = gdf[gdf['FIRST_DNAM'] == district_name].copy()
    
    if len(district_gdf) == 0:
        st.error(f"No chiefdoms found for {district_name} district in shapefile")
        return None
    
    # Get unique chiefdoms from shapefile
    chiefdoms = sorted(district_gdf['FIRST_CHIE'].dropna().unique())
    
    # Calculate rows needed
    rows = math.ceil(len(chiefdoms) / cols)
    
    # Create subplot figure with increased vertical space
    fig, axes = plt.subplots(rows, cols, figsize=(cols*5, rows*6))
    fig.suptitle(f'{district_name} District - ITN Coverage Analysis', 
                 fontsize=20, fontweight='bold', y=0.98)
    
    # Ensure axes is always 2D array
    if rows == 1:
        axes = axes.reshape(1, -1)
    elif cols == 1:
        axes = axes.reshape(-1, 1)
    
    # Plot each chiefdom
    for idx, chiefdom in enumerate(chiefdoms):
        row = idx // cols
        col = idx % cols
        ax = axes[row, col]
        
        # Filter shapefile for this specific chiefdom
        chiefdom_gdf = district_gdf[district_gdf['FIRST_CHIE'] == chiefdom].copy()
        
        # Filter ITN data for this district and chiefdom
        district_data = itn_df[itn_df["District"].str.upper() == district_name.upper()].copy()
        chiefdom_data = district_data[district_data["Chiefdom"] == chiefdom].copy()
        
        # Calculate totals for this chiefdom
        enrollment_total = int(chiefdom_data["Total_Enrollment"].sum()) if len(chiefdom_data) > 0 else 0
        itns_total = int(chiefdom_data["Total_ITNs"].sum()) if len(chiefdom_data) > 0 else 0
        
        # Calculate coverage percentage
        coverage_percent = (itns_total / enrollment_total * 100) if enrollment_total > 0 else 0
        coverage_percent = min(coverage_percent, 100)  # Cap at 100%
        
        # Get color based on coverage (same as Section 2)
        coverage_color = get_coverage_color(coverage_percent)
        
        # Plot chiefdom boundary with coverage color
        chiefdom_gdf.plot(ax=ax, color=coverage_color, edgecolor='black', alpha=0.8, linewidth=2)
        
        # Create ITN coverage text (n, m) format
        itn_text = f"({itns_total}, {enrollment_total})"
        
        # Set title with ITN coverage information
        ax.set_title(f'{chiefdom}\n{itn_text}', 
                    fontsize=11, fontweight='bold', pad=10)
        
        # Add coverage percentage in the center of the chiefdom
        if len(chiefdom_gdf) > 0:
            # Get center of chiefdom
            bounds = chiefdom_gdf.total_bounds
            center_x = (bounds[0] + bounds[2]) / 2
            center_y = (bounds[1] + bounds[3]) / 2
            
            # Add coverage percentage text in the center
            ax.text(center_x, center_y, f"{coverage_percent:.0f}%", 
                   fontsize=16, fontweight='bold', color='white', 
                   ha='center', va='center',
                   bbox=dict(boxstyle='round,pad=0.5', facecolor='black', alpha=0.7))
        
        # Remove axis labels and ticks for cleaner look
        ax.set_xticks([])
        ax.set_yticks([])
        ax.set_xlabel('')
        ax.set_ylabel('')
        
        # Remove the box frame
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.spines['bottom'].set_visible(False)
        ax.spines['left'].set_visible(False)
        
        # Add light grid
        ax.grid(True, alpha=0.2, linestyle='--')
        
        # Set equal aspect ratio and tight layout
        ax.set_aspect('equal')
        
        # Set bounds to chiefdom extent with some padding
        bounds = chiefdom_gdf.total_bounds
        padding = 0.01
        ax.set_xlim(bounds[0] - padding, bounds[2] + padding)
        ax.set_ylim(bounds[1] - padding, bounds[3] + padding)
    
    # Hide empty subplots
    total_plots = rows * cols
    for idx in range(len(chiefdoms), total_plots):
        row = idx // cols
        col = idx % cols
        axes[row, col].set_visible(False)
    
    plt.tight_layout()
    plt.subplots_adjust(top=0.93, hspace=0.4, wspace=0.3)
    
    return fig

# Extract ITN coverage data
try:
    itn_df = extract_itn_data_from_excel(df_original)
except Exception as e:
    st.error(f"Error extracting ITN data: {e}")
    itn_df = pd.DataFrame()

# BO District ITN Coverage Dashboard
st.subheader("3a. BO District - ITN Coverage")

with st.spinner("Generating BO District ITN coverage dashboard..."):
    try:
        fig_bo_itn = create_itn_coverage_dashboard(gdf, itn_df, "BO", columns)
        if fig_bo_itn:
            st.pyplot(fig_bo_itn)
            
            # Save figure option
            buffer_bo_itn = BytesIO()
            fig_bo_itn.savefig(buffer_bo_itn, format='png', dpi=300, bbox_inches='tight')
            buffer_bo_itn.seek(0)
            
            st.download_button(
                label="📥 Download BO District ITN Coverage Dashboard",
                data=buffer_bo_itn,
                file_name="BO_District_ITN_Coverage_Dashboard.png",
                mime="image/png"
            )
        else:
            st.warning("Could not generate BO District ITN coverage dashboard")
    except Exception as e:
        st.error(f"Error generating BO District ITN coverage dashboard: {e}")

st.divider()

# BOMBALI District ITN Coverage Dashboard
st.subheader("3b. BOMBALI District - ITN Coverage")

with st.spinner("Generating BOMBALI District ITN coverage dashboard..."):
    try:
        fig_bombali_itn = create_itn_coverage_dashboard(gdf, itn_df, "BOMBALI", columns)
        if fig_bombali_itn:
            st.pyplot(fig_bombali_itn)
            
            # Save figure option
            buffer_bombali_itn = BytesIO()
            fig_bombali_itn.savefig(buffer_bombali_itn, format='png', dpi=300, bbox_inches='tight')
            buffer_bombali_itn.seek(0)
            
            st.download_button(
                label="📥 Download BOMBALI District ITN Coverage Dashboard",
                data=buffer_bombali_itn,
                file_name="BOMBALI_District_ITN_Coverage_Dashboard.png",
                mime="image/png"
            )
        else:
            st.warning("Could not generate BOMBALI District ITN coverage dashboard")
    except Exception as e:
        st.error(f"Error generating BOMBALI District ITN coverage dashboard: {e}")

# Section 4: Summary of Key Findings
st.header("📋 Section 4: Summary of Key Findings")

def generate_key_summary(extracted_df, itn_df):
    """Generate summary of key findings across all sections"""
    
    summary = {}
    
    # Overall statistics
    total_schools = len(extracted_df)
    total_districts = len(extracted_df['District'].dropna().unique())
    total_gps_records = len(extracted_df[extracted_df['GPS_Location'].notna()])
    
    # District breakdown
    bo_schools = len(extracted_df[extracted_df['District'].str.upper() == 'BO'])
    bombali_schools = len(extracted_df[extracted_df['District'].str.upper() == 'BOMBALI'])
    
    # GPS coverage
    gps_coverage = (total_gps_records / total_schools * 100) if total_schools > 0 else 0
    
    # ITN statistics from itn_df
    total_enrollment = int(itn_df['Total_Enrollment'].sum())
    total_itns_distributed = int(itn_df['Distributed_ITNs'].sum())
    total_itns_available = int(itn_df['Total_ITNs'].sum())
    
    # Coverage calculations
    enrollment_coverage = (total_itns_distributed / total_enrollment * 100) if total_enrollment > 0 else 0
    distribution_efficiency = (total_itns_distributed / total_itns_available * 100) if total_itns_available > 0 else 0
    
    # School coverage analysis (using real target data)
    target_data = generate_target_school_data([])  # Get all target data
    total_target_schools = sum(target_data.values())
    school_coverage = (total_schools / total_target_schools * 100) if total_target_schools > 0 else 0
    
    summary = {
        'total_schools': total_schools,
        'total_districts': total_districts,
        'bo_schools': bo_schools,
        'bombali_schools': bombali_schools,
        'gps_coverage': gps_coverage,
        'total_enrollment': total_enrollment,
        'total_itns_distributed': total_itns_distributed,
        'total_itns_available': total_itns_available,
        'enrollment_coverage': enrollment_coverage,
        'distribution_efficiency': distribution_efficiency,
        'total_target_schools': total_target_schools,
        'school_coverage': school_coverage
    }
    
    return summary

# Generate and display summary
try:
    key_summary = generate_key_summary(extracted_df, itn_df)
    
    # Key Metrics Display
    st.subheader("📊 Key Metrics Overview")
    
    # Row 1: Basic counts
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Total Schools Surveyed", f"{key_summary['total_schools']:,}")
    with col2:
        st.metric("Districts Covered", f"{key_summary['total_districts']}")
    with col3:
        st.metric("BO District Schools", f"{key_summary['bo_schools']:,}")
    with col4:
        st.metric("BOMBALI District Schools", f"{key_summary['bombali_schools']:,}")
    
    # Row 2: Coverage metrics
    col5, col6, col7, col8 = st.columns(4)
    with col5:
        st.metric("GPS Coverage", f"{key_summary['gps_coverage']:.1f}%")
    with col6:
        st.metric("School Survey Coverage", f"{key_summary['school_coverage']:.1f}%")
    with col7:
        st.metric("ITN Enrollment Coverage", f"{key_summary['enrollment_coverage']:.1f}%")
    with col8:
        st.metric("ITN Distribution Efficiency", f"{key_summary['distribution_efficiency']:.1f}%")
    
    # Row 3: ITN statistics
    col9, col10, col11, col12 = st.columns(4)
    with col9:
        st.metric("Total Student Enrollment", f"{key_summary['total_enrollment']:,}")
    with col10:
        st.metric("ITNs Distributed", f"{key_summary['total_itns_distributed']:,}")
    with col11:
        st.metric("Total ITNs Available", f"{key_summary['total_itns_available']:,}")
    with col12:
        st.metric("Target Schools", f"{key_summary['total_target_schools']:,}")
    
    # Summary Insights
    st.subheader("🔍 Key Insights")
    
    insights_col1, insights_col2 = st.columns(2)
    
    with insights_col1:
        st.markdown("### 📍 Geographic Coverage")
        st.write(f"• **{key_summary['gps_coverage']:.1f}%** of schools have GPS coordinates")
        st.write(f"• **{key_summary['school_coverage']:.1f}%** of target schools surveyed")
        st.write(f"• **{key_summary['total_schools']:,}** schools across **{key_summary['total_districts']}** districts")
        
        if key_summary['gps_coverage'] >= 80:
            st.success("✅ Excellent GPS coverage")
        elif key_summary['gps_coverage'] >= 60:
            st.warning("⚠️ Good GPS coverage")
        else:
            st.error("❌ Poor GPS coverage - needs improvement")
    
    with insights_col2:
        st.markdown("### 🛡️ ITN Distribution")
        st.write(f"• **{key_summary['enrollment_coverage']:.1f}%** enrollment coverage")
        st.write(f"• **{key_summary['distribution_efficiency']:.1f}%** distribution efficiency")
        st.write(f"• **{key_summary['total_itns_distributed']:,}** ITNs distributed to students")
        
        if key_summary['enrollment_coverage'] >= 80:
            st.success("✅ Excellent ITN coverage")
        elif key_summary['enrollment_coverage'] >= 60:
            st.warning("⚠️ Good ITN coverage")
        else:
            st.error("❌ Poor ITN coverage - needs improvement")
    
    # District Comparison
    st.subheader("⚖️ District Comparison")
    
    # Calculate district-specific metrics
    bo_data = extracted_df[extracted_df['District'].str.upper() == 'BO']
    bombali_data = extracted_df[extracted_df['District'].str.upper() == 'BOMBALI']
    
    bo_gps_coverage = (len(bo_data[bo_data['GPS_Location'].notna()]) / len(bo_data) * 100) if len(bo_data) > 0 else 0
    bombali_gps_coverage = (len(bombali_data[bombali_data['GPS_Location'].notna()]) / len(bombali_data) * 100) if len(bombali_data) > 0 else 0
    
    # District comparison table
    bo_chiefdoms = ['BADJIA', 'BAGBWE(BAGBE)', 'BOAMA', 'BAGBO', 'BO TOWN', 'BONGOR', 'BUMPE NGAO', 'GBO', 'JAIAMA', 'KAKUA', 'KOMBOYA', 'LUGBU', 'NIAWA LENGA', 'SELENGA', 'TIKONKO', 'VALUNIA', 'WONDE']
    bombali_chiefdoms = ['BIRIWA', 'BOMBALI SEBORA', 'BOMBALI SIARI', 'GBANTI', 'GBENDEMBU', 'KAMARANKA', 'MAGBAIMBA NDORWAHUN', 'MAKARI', 'MAKENI CITY', 'MARA', 'N\'GOWAHUN', 'PAKI MASABONG', 'SAFROKO LIMBA']
    
    target_data_all = generate_target_school_data([])
    bo_target_total = sum([v for k, v in target_data_all.items() if k in bo_chiefdoms])
    bombali_target_total = sum([v for k, v in target_data_all.items() if k in bombali_chiefdoms])
    
    comparison_data = {
        'Metric': ['Schools Surveyed', 'GPS Coverage (%)', 'Target Schools', 'Survey Coverage (%)'],
        'BO District': [
            f"{key_summary['bo_schools']:,}",
            f"{bo_gps_coverage:.1f}%", 
            f"{bo_target_total:,}",
            f"{(key_summary['bo_schools'] / bo_target_total * 100):.1f}%" if bo_target_total > 0 else "0.0%"
        ],
        'BOMBALI District': [
            f"{key_summary['bombali_schools']:,}",
            f"{bombali_gps_coverage:.1f}%",
            f"{bombali_target_total:,}",
            f"{(key_summary['bombali_schools'] / bombali_target_total * 100):.1f}%" if bombali_target_total > 0 else "0.0%"
        ]
    }
    
    comparison_df = pd.DataFrame(comparison_data)
    st.dataframe(comparison_df, use_container_width=True)
    
    # Action Items
    st.subheader("🎯 Recommended Actions")
    
    action_col1, action_col2 = st.columns(2)
    
    with action_col1:
        st.markdown("### 📈 Improve Coverage")
        if key_summary['school_coverage'] < 80:
            st.write("• Increase school survey efforts in under-covered chiefdoms")
        if key_summary['gps_coverage'] < 90:
            st.write("• Collect GPS coordinates for remaining schools")
        st.write("• Focus on chiefdoms with red/orange coverage indicators")
        st.write("• Prioritize high-enrollment areas")
    
    with action_col2:
        st.markdown("### 🛡️ Enhance ITN Distribution")
        if key_summary['enrollment_coverage'] < 80:
            st.write("• Increase ITN distribution in low-coverage areas")
        if key_summary['distribution_efficiency'] < 90:
            st.write("• Improve distribution of available ITNs")
        st.write("• Target schools with high enrollment but low ITN coverage")
        st.write("• Ensure absent students receive ITNs")
    
except Exception as e:
    st.error(f"Error generating summary: {e}")

# Section 5: Summary of Distribution and Coverage
st.header("📊 Section 5: Summary of Distribution and Coverage")

def generate_district_chiefdom_analysis(extracted_df, itn_df, gdf):
    """Generate comprehensive district and chiefdom analysis"""
    
    analysis = {}
    
    for district_name in ["BO", "BOMBALI"]:
        # Filter data for this district
        district_data = extracted_df[extracted_df["District"].str.upper() == district_name.upper()].copy()
        district_itn_data = itn_df[itn_df["District"].str.upper() == district_name.upper()].copy()
        district_gdf = gdf[gdf['FIRST_DNAM'] == district_name].copy()
        
        # Get chiefdoms for this district
        chiefdoms = sorted(district_gdf['FIRST_CHIE'].dropna().unique())
        
        chiefdom_analysis = []
        
        for chiefdom in chiefdoms:
            # Filter data for this chiefdom
            chiefdom_data = district_data[district_data["Chiefdom"] == chiefdom].copy()
            chiefdom_itn_data = district_itn_data[district_itn_data["Chiefdom"] == chiefdom].copy()
            
            # Get target schools
            target_data = generate_target_school_data([])
            target_schools = target_data.get(chiefdom, 0)
            
            # Calculate metrics
            schools_surveyed = len(chiefdom_data)
            gps_schools = len(chiefdom_data[chiefdom_data['GPS_Location'].notna()])
            total_enrollment = int(chiefdom_itn_data["Total_Enrollment"].sum()) if len(chiefdom_itn_data) > 0 else 0
            itns_distributed = int(chiefdom_itn_data["Distributed_ITNs"].sum()) if len(chiefdom_itn_data) > 0 else 0
            total_itns = int(chiefdom_itn_data["Total_ITNs"].sum()) if len(chiefdom_itn_data) > 0 else 0
            
            # Calculate percentages
            school_coverage = (schools_surveyed / target_schools * 100) if target_schools > 0 else 0
            gps_coverage = (gps_schools / schools_surveyed * 100) if schools_surveyed > 0 else 0
            itn_coverage = (itns_distributed / total_enrollment * 100) if total_enrollment > 0 else 0
            
            chiefdom_analysis.append({
                'Chiefdom': chiefdom,
                'Target_Schools': target_schools,
                'Schools_Surveyed': schools_surveyed,
                'GPS_Schools': gps_schools,
                'Total_Enrollment': total_enrollment,
                'ITNs_Distributed': itns_distributed,
                'Total_ITNs': total_itns,
                'School_Coverage': school_coverage,
                'GPS_Coverage': gps_coverage,
                'ITN_Coverage': itn_coverage
            })
        
        analysis[district_name] = chiefdom_analysis
    
    return analysis

# Generate analysis
try:
    district_analysis = generate_district_chiefdom_analysis(extracted_df, itn_df, gdf)
    
    # District Comparison Overview
    st.subheader("🏘️ District Comparison Overview")
    
    # Calculate district totals
    district_totals = {}
    for district in ["BO", "BOMBALI"]:
        district_data = district_analysis[district]
        
        total_target = sum([d['Target_Schools'] for d in district_data])
        total_surveyed = sum([d['Schools_Surveyed'] for d in district_data])
        total_gps = sum([d['GPS_Schools'] for d in district_data])
        total_enrollment = sum([d['Total_Enrollment'] for d in district_data])
        total_itns_dist = sum([d['ITNs_Distributed'] for d in district_data])
        total_itns = sum([d['Total_ITNs'] for d in district_data])
        
        district_totals[district] = {
            'Target_Schools': total_target,
            'Schools_Surveyed': total_surveyed,
            'GPS_Schools': total_gps,
            'Total_Enrollment': total_enrollment,
            'ITNs_Distributed': total_itns_dist,
            'Total_ITNs': total_itns,
            'School_Coverage': (total_surveyed / total_target * 100) if total_target > 0 else 0,
            'GPS_Coverage': (total_gps / total_surveyed * 100) if total_surveyed > 0 else 0,
            'ITN_Coverage': (total_itns_dist / total_enrollment * 100) if total_enrollment > 0 else 0
        }
    
    # Display district comparison
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### BO District Summary")
        bo_data = district_totals["BO"]
        st.metric("Target Schools", f"{bo_data['Target_Schools']:,}")
        st.metric("Schools Surveyed", f"{bo_data['Schools_Surveyed']:,}")
        st.metric("School Coverage", f"{bo_data['School_Coverage']:.1f}%")
        st.metric("GPS Coverage", f"{bo_data['GPS_Coverage']:.1f}%")
        st.metric("Total Enrollment", f"{bo_data['Total_Enrollment']:,}")
        st.metric("ITNs Distributed", f"{bo_data['ITNs_Distributed']:,}")
        st.metric("ITN Coverage", f"{bo_data['ITN_Coverage']:.1f}%")
    
    with col2:
        st.markdown("### BOMBALI District Summary")
        bombali_data = district_totals["BOMBALI"]
        st.metric("Target Schools", f"{bombali_data['Target_Schools']:,}")
        st.metric("Schools Surveyed", f"{bombali_data['Schools_Surveyed']:,}")
        st.metric("School Coverage", f"{bombali_data['School_Coverage']:.1f}%")
        st.metric("GPS Coverage", f"{bombali_data['GPS_Coverage']:.1f}%")
        st.metric("Total Enrollment", f"{bombali_data['Total_Enrollment']:,}")
        st.metric("ITNs Distributed", f"{bombali_data['ITNs_Distributed']:,}")
        st.metric("ITN Coverage", f"{bombali_data['ITN_Coverage']:.1f}%")
    
    # Chiefdom Analysis by District
    st.subheader("🏘️ Chiefdom Analysis by District")
    
    # BO District Chiefdom Analysis
    st.markdown("### BO District - Chiefdom Performance")
    bo_df = pd.DataFrame(district_analysis["BO"])
    bo_df = bo_df.sort_values('ITN_Coverage', ascending=False)
    
    # Format percentages for display
    bo_display = bo_df.copy()
    bo_display['School_Coverage'] = bo_display['School_Coverage'].apply(lambda x: f"{x:.1f}%")
    bo_display['GPS_Coverage'] = bo_display['GPS_Coverage'].apply(lambda x: f"{x:.1f}%")
    bo_display['ITN_Coverage'] = bo_display['ITN_Coverage'].apply(lambda x: f"{x:.1f}%")
    
    st.dataframe(bo_display, use_container_width=True)
    
    # BOMBALI District Chiefdom Analysis
    st.markdown("### BOMBALI District - Chiefdom Performance")
    bombali_df = pd.DataFrame(district_analysis["BOMBALI"])
    bombali_df = bombali_df.sort_values('ITN_Coverage', ascending=False)
    
    # Format percentages for display
    bombali_display = bombali_df.copy()
    bombali_display['School_Coverage'] = bombali_display['School_Coverage'].apply(lambda x: f"{x:.1f}%")
    bombali_display['GPS_Coverage'] = bombali_display['GPS_Coverage'].apply(lambda x: f"{x:.1f}%")
    bombali_display['ITN_Coverage'] = bombali_display['ITN_Coverage'].apply(lambda x: f"{x:.1f}%")
    
    st.dataframe(bombali_display, use_container_width=True)
    
    # Performance Rankings
    st.subheader("🏆 Performance Rankings")
    
    # Combine all chiefdoms for ranking
    all_chiefdoms = []
    for district in ["BO", "BOMBALI"]:
        for chiefdom_data in district_analysis[district]:
            chiefdom_data['District'] = district
            all_chiefdoms.append(chiefdom_data)
    
    all_chiefdoms_df = pd.DataFrame(all_chiefdoms)
    
    ranking_col1, ranking_col2, ranking_col3 = st.columns(3)
    
    with ranking_col1:
        st.markdown("#### 🥇 Top ITN Coverage")
        top_itn = all_chiefdoms_df.nlargest(5, 'ITN_Coverage')[['District', 'Chiefdom', 'ITN_Coverage']]
        top_itn['ITN_Coverage'] = top_itn['ITN_Coverage'].apply(lambda x: f"{x:.1f}%")
        st.dataframe(top_itn, hide_index=True)
    
    with ranking_col2:
        st.markdown("#### 📊 Top School Coverage")
        top_school = all_chiefdoms_df.nlargest(5, 'School_Coverage')[['District', 'Chiefdom', 'School_Coverage']]
        top_school['School_Coverage'] = top_school['School_Coverage'].apply(lambda x: f"{x:.1f}%")
        st.dataframe(top_school, hide_index=True)
    
    with ranking_col3:
        st.markdown("#### 📍 Top GPS Coverage")
        top_gps = all_chiefdoms_df.nlargest(5, 'GPS_Coverage')[['District', 'Chiefdom', 'GPS_Coverage']]
        top_gps['GPS_Coverage'] = top_gps['GPS_Coverage'].apply(lambda x: f"{x:.1f}%")
        st.dataframe(top_gps, hide_index=True)
    
    # Areas Needing Attention
    st.subheader("⚠️ Areas Needing Attention")
    
    attention_col1, attention_col2 = st.columns(2)
    
    with attention_col1:
        st.markdown("#### 🔴 Low ITN Coverage (< 50%)")
        low_itn = all_chiefdoms_df[all_chiefdoms_df['ITN_Coverage'] < 50][['District', 'Chiefdom', 'ITN_Coverage', 'Total_Enrollment']]
        if len(low_itn) > 0:
            low_itn = low_itn.sort_values('ITN_Coverage')
            low_itn['ITN_Coverage'] = low_itn['ITN_Coverage'].apply(lambda x: f"{x:.1f}%")
            st.dataframe(low_itn, hide_index=True)
        else:
            st.success("✅ All chiefdoms have ITN coverage ≥ 50%")
    
    with attention_col2:
        st.markdown("#### 🔴 Low School Coverage (< 50%)")
        low_school = all_chiefdoms_df[all_chiefdoms_df['School_Coverage'] < 50][['District', 'Chiefdom', 'School_Coverage', 'Target_Schools']]
        if len(low_school) > 0:
            low_school = low_school.sort_values('School_Coverage')
            low_school['School_Coverage'] = low_school['School_Coverage'].apply(lambda x: f"{x:.1f}%")
            st.dataframe(low_school, hide_index=True)
        else:
            st.success("✅ All chiefdoms have school coverage ≥ 50%")

except Exception as e:
    st.error(f"Error generating distribution and coverage analysis: {e}")

# Raw data preview
if st.checkbox("Show raw data preview"):
    st.subheader("📄 Raw Data Preview")
    st.dataframe(extracted_df.head(20))
    
    # Download raw data
    csv_data = extracted_df.to_csv(index=False)
    st.download_button(
        label="📥 Download Extracted Data as CSV",
        data=csv_data,
        file_name="extracted_gps_data.csv",
        mime="text/csv"
    )

# Final Export Section
st.header("📥 Export Complete Dashboard Reports")
st.write("Generate comprehensive reports with all sections and visualizations")

# Current date for reports
from datetime import datetime
current_date = datetime.now()
date_str = current_date.strftime("%B %d, %Y")
filename_date = current_date.strftime("%Y%m%d_%H%M")

# Create columns for export buttons
export_col1, export_col2 = st.columns(2)

with export_col1:
    # PDF Report Generation
    if st.button("📋 Generate PDF Report", help="Generate comprehensive PDF report with all dashboards"):
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, PageBreak, Table, TableStyle
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib import colors
            from reportlab.lib.enums import TA_CENTER, TA_LEFT
            import tempfile
            import os
            
            # Create PDF buffer
            pdf_buffer = BytesIO()
            
            # Create PDF document
            doc = SimpleDocTemplate(pdf_buffer, pagesize=A4, topMargin=0.5*inch, bottomMargin=0.5*inch)
            story = []
            
            # Get styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                alignment=TA_CENTER,
                spaceAfter=30
            )
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=16,
                alignment=TA_LEFT,
                spaceAfter=12
            )
            normal_style = styles['Normal']
            
            # Header with logos from file paths
            logo_config = st.session_state.get('logos', {
                'logo1': {'path': 'logo1.png', 'name': 'Organization 1'},
                'logo2': {'path': 'logo2.png', 'name': 'Organization 2'},
                'logo3': {'path': 'logo3.png', 'name': 'Organization 3'},
                'logo4': {'path': 'logo4.png', 'name': 'Organization 4'}
            })
            
            # Try to add actual logos or use placeholders
            logo_row_data = []
            name_row_data = []
            
            for i in range(1, 5):
                logo_info = logo_config.get(f'logo{i}', {'path': f'logo{i}.png', 'name': f'Organization {i}'})
                try:
                    # Try to add image (placeholder for now in PDF)
                    logo_row_data.append(f'[LOGO {i}]')
                except:
                    logo_row_data.append(f'[LOGO {i}]')
                name_row_data.append(logo_info['name'])
            
            header_table_data = [logo_row_data, name_row_data]
            header_table = Table(header_table_data, colWidths=[2*inch, 2*inch, 2*inch, 2*inch])
            header_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('FONTSIZE', (0, 1), (-1, 1), 10),
                ('FONTNAME', (0, 1), (-1, 1), 'Helvetica-Bold'),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('BACKGROUND', (0, 0), (-1, -1), colors.lightgrey),
            ]))
            story.append(header_table)
            story.append(Spacer(1, 20))
            
            # Title Page
            story.append(Paragraph("School-Based Distribution (SBD)", title_style))
            story.append(Paragraph("Comprehensive Chiefdom Dashboard Analysis", heading_style))
            story.append(Spacer(1, 20))
            
            # Date
            story.append(Paragraph(f"Report Generated: {date_str}", normal_style))
            story.append(Spacer(1, 40))
            
            # Executive Summary
            summary_text = f"""
            <b>EXECUTIVE SUMMARY</b><br/><br/>
            This comprehensive dashboard report presents analysis across three key areas for BO and BOMBALI districts:<br/><br/>
            
            <b>Section 1: GPS School Locations</b><br/>
            • Visual mapping of all school GPS coordinates by chiefdom<br/>
            • Geographic distribution analysis across {len(gdf[gdf['FIRST_DNAM'].isin(['BO', 'BOMBALI'])])} chiefdoms<br/><br/>
            
            <b>Section 2: School Coverage Analysis</b><br/>
            • Comparison of surveyed schools vs target schools per chiefdom<br/>
            • Coverage rates with color-coded visualization<br/>
            • Identification of gaps in school survey completion<br/><br/>
            
            <b>Section 3: ITN Coverage Analysis</b><br/>
            • ITN distribution effectiveness by chiefdom<br/>
            • Coverage percentages based on enrollment vs ITN distribution<br/>
            • Performance assessment across administrative boundaries<br/><br/>
            
            <b>Districts Covered:</b> BO, BOMBALI<br/>
            <b>Total Records Processed:</b> {len(extracted_df)}<br/>
            <b>Report Date:</b> {date_str}
            """
            story.append(Paragraph(summary_text, normal_style))
            story.append(PageBreak())
            
            # Dashboard Sections
            sections = [
                "Section 1: GPS School Locations - Visual mapping of school coordinates",
                "Section 2: School Coverage Analysis - Survey completion rates",
                "Section 3: ITN Coverage Analysis - Distribution effectiveness"
            ]
            
            for section in sections:
                story.append(Paragraph(section, heading_style))
                story.append(Paragraph("Detailed visualizations showing chiefdom-level analysis with color-coded performance indicators.", normal_style))
                story.append(Spacer(1, 20))
            
            # Data Summary Table
            story.append(Paragraph("Data Summary", heading_style))
            summary_table_data = [
                ['District', 'Chiefdoms', 'Records', 'GPS Records'],
                ['BO', str(len(gdf[gdf['FIRST_DNAM'] == 'BO'])), 
                 str(len(extracted_df[extracted_df['District'].str.upper() == 'BO'])),
                 str(len(extracted_df[(extracted_df['District'].str.upper() == 'BO') & extracted_df['GPS_Location'].notna()]))],
                ['BOMBALI', str(len(gdf[gdf['FIRST_DNAM'] == 'BOMBALI'])), 
                 str(len(extracted_df[extracted_df['District'].str.upper() == 'BOMBALI'])),
                 str(len(extracted_df[(extracted_df['District'].str.upper() == 'BOMBALI') & extracted_df['GPS_Location'].notna()]))]
            ]
            
            summary_table = Table(summary_table_data, colWidths=[2*inch, 1.5*inch, 1.5*inch, 1.5*inch])
            summary_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(summary_table)
            
            # Build PDF
            doc.build(story)
            
            # Get PDF data
            pdf_data = pdf_buffer.getvalue()
            
            st.success("✅ PDF report generated successfully!")
            st.download_button(
                label="💾 Download PDF Report",
                data=pdf_data,
                file_name=f"SBD_Dashboard_Report_{filename_date}.pdf",
                mime="application/pdf",
                help="Download comprehensive PDF report with all dashboard sections"
            )
            
        except ImportError:
            st.error("❌ PDF generation requires reportlab library. Please install it using: pip install reportlab")
        except Exception as e:
            st.error(f"❌ Error generating PDF: {str(e)}")

with export_col2:
    # Word Report Generation
    if st.button("📋 Generate Word Report", help="Generate comprehensive Word report with all dashboards"):
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.oxml.shared import OxmlElement, qn
            
            doc = Document()
            
            # Header with logos from file paths
            logo_config = st.session_state.get('logos', {
                'logo1': {'path': 'logo1.png', 'name': 'Organization 1'},
                'logo2': {'path': 'logo2.png', 'name': 'Organization 2'},
                'logo3': {'path': 'logo3.png', 'name': 'Organization 3'},
                'logo4': {'path': 'logo4.png', 'name': 'Organization 4'}
            })
            
            # Header with logo placeholders and organization names
            header_table = doc.add_table(rows=2, cols=4)
            header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Logo placeholders row
            logo_cells = header_table.rows[0].cells
            org_cells = header_table.rows[1].cells
            
            for i in range(4):
                logo_info = logo_config.get(f'logo{i+1}', {'path': f'logo{i+1}.png', 'name': f'Organization {i+1}'})
                
                # Try to add actual logo or use placeholder
                try:
                    # Attempt to add image from file path
                    logo_para = logo_cells[i].paragraphs[0]
                    logo_run = logo_para.add_run()
                    logo_run.add_picture(logo_info['path'], width=Inches(1.5))
                    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except:
                    # Use placeholder if file not found
                    logo_cells[i].text = f"[LOGO {i+1}]"
                
                # Add organization name
                org_cells[i].text = logo_info['name']
            
            # Style header table
            for row in header_table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = cell.paragraphs[0].runs[0]
                    run.font.size = Pt(10)
                    run.bold = True
            
            doc.add_paragraph()  # Add space after header
            
            # Title
            title = doc.add_heading('School-Based Distribution (SBD)', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle = doc.add_heading('Comprehensive Chiefdom Dashboard Analysis', level=1)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Date
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_run = date_para.add_run(f"Report Generated: {date_str}")
            date_run.font.size = Pt(12)
            date_run.bold = True
            
            doc.add_page_break()
            
            # Executive Summary
            doc.add_heading('Executive Summary', level=1)
            
            summary_text = f"""
            This comprehensive dashboard report presents analysis across three key areas for BO and BOMBALI districts:
            
            Section 1: GPS School Locations
            • Visual mapping of all school GPS coordinates by chiefdom
            • Geographic distribution analysis across {len(gdf[gdf['FIRST_DNAM'].isin(['BO', 'BOMBALI'])])} chiefdoms
            
            Section 2: School Coverage Analysis
            • Comparison of surveyed schools vs target schools per chiefdom
            • Coverage rates with color-coded visualization
            • Identification of gaps in school survey completion
            
            Section 3: ITN Coverage Analysis
            • ITN distribution effectiveness by chiefdom
            • Coverage percentages based on enrollment vs ITN distribution
            • Performance assessment across administrative boundaries
            
            Districts Covered: BO, BOMBALI
            Total Records Processed: {len(extracted_df)}
            Report Date: {date_str}
            """
            doc.add_paragraph(summary_text)
            
            # Dashboard Sections Overview
            doc.add_heading('Dashboard Sections', level=1)
            
            sections = [
                ("Section 1: GPS School Locations", "Visual mapping of school coordinates by chiefdom with geographic distribution analysis"),
                ("Section 2: School Coverage Analysis", "Survey completion rates comparing actual vs target schools with color-coded performance"),
                ("Section 3: ITN Coverage Analysis", "ITN distribution effectiveness showing coverage percentages and performance indicators")
            ]
            
            for section_title, section_desc in sections:
                doc.add_heading(section_title, level=2)
                doc.add_paragraph(section_desc)
            
            # Data Summary
            doc.add_heading('Data Summary', level=1)
            
            # Create summary table
            summary_table = doc.add_table(rows=3, cols=4)
            summary_table.style = 'Table Grid'
            
            # Header row
            header_cells = summary_table.rows[0].cells
            header_cells[0].text = 'District'
            header_cells[1].text = 'Chiefdoms'
            header_cells[2].text = 'Records'
            header_cells[3].text = 'GPS Records'
            
            # BO row
            bo_cells = summary_table.rows[1].cells
            bo_cells[0].text = 'BO'
            bo_cells[1].text = str(len(gdf[gdf['FIRST_DNAM'] == 'BO']))
            bo_cells[2].text = str(len(extracted_df[extracted_df['District'].str.upper() == 'BO']))
            bo_cells[3].text = str(len(extracted_df[(extracted_df['District'].str.upper() == 'BO') & extracted_df['GPS_Location'].notna()]))
            
            # BOMBALI row
            bombali_cells = summary_table.rows[2].cells
            bombali_cells[0].text = 'BOMBALI'
            bombali_cells[1].text = str(len(gdf[gdf['FIRST_DNAM'] == 'BOMBALI']))
            bombali_cells[2].text = str(len(extracted_df[extracted_df['District'].str.upper() == 'BOMBALI']))
            bombali_cells[3].text = str(len(extracted_df[(extracted_df['District'].str.upper() == 'BOMBALI') & extracted_df['GPS_Location'].notna()]))
            
            # Style summary table
            for row in summary_table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Bold header row
            for cell in summary_table.rows[0].cells:
                run = cell.paragraphs[0].runs[0]
                run.bold = True
            
            # Save to BytesIO
            word_buffer = BytesIO()
            doc.save(word_buffer)
            word_data = word_buffer.getvalue()
            
            st.success("✅ Word report generated successfully!")
            st.download_button(
                label="💾 Download Word Report",
                data=word_data,
                file_name=f"SBD_Dashboard_Report_{filename_date}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                help="Download comprehensive Word report with all dashboard sections"
            )
            
        except ImportError:
            st.error("❌ Word report generation requires python-docx library. Please install it using: pip install python-docx")
        except Exception as e:
            st.error(f"❌ Error generating Word report: {str(e)}")

# Logo Management Section
st.header("🖼️ Logo Management")
st.write("Configure logo file paths for report headers")

logo_col1, logo_col2, logo_col3, logo_col4 = st.columns(4)

with logo_col1:
    st.write("**Logo 1 (Organization 1)**")
    logo1_path = st.text_input("Logo 1 file path", value="logo1.png", key="logo1_path")
    org1_name = st.text_input("Organization 1 name", value="Organization 1", key="org1_name")
    try:
        st.image(logo1_path, width=150, caption="Logo 1 Preview")
    except:
        st.write("❌ Logo 1 file not found")

with logo_col2:
    st.write("**Logo 2 (Organization 2)**")
    logo2_path = st.text_input("Logo 2 file path", value="logo2.png", key="logo2_path")
    org2_name = st.text_input("Organization 2 name", value="Organization 2", key="org2_name")
    try:
        st.image(logo2_path, width=150, caption="Logo 2 Preview")
    except:
        st.write("❌ Logo 2 file not found")

with logo_col3:
    st.write("**Logo 3 (Organization 3)**")
    logo3_path = st.text_input("Logo 3 file path", value="logo3.png", key="logo3_path")
    org3_name = st.text_input("Organization 3 name", value="Organization 3", key="org3_name")
    try:
        st.image(logo3_path, width=150, caption="Logo 3 Preview")
    except:
        st.write("❌ Logo 3 file not found")

with logo_col4:
    st.write("**Logo 4 (Organization 4)**")
    logo4_path = st.text_input("Logo 4 file path", value="logo4.png", key="logo4_path")
    org4_name = st.text_input("Organization 4 name", value="Organization 4", key="org4_name")
    try:
        st.image(logo4_path, width=150, caption="Logo 4 Preview")
    except:
        st.write("❌ Logo 4 file not found")


# Update logo configuration in session state
st.session_state.logos = {
    'logo1': {'path': logo1_path, 'name': org1_name},
    'logo2': {'path': logo2_path, 'name': org2_name},
    'logo3': {'path': logo3_path, 'name': org3_name},
    'logo4': {'path': logo4_path, 'name': org4_name}
}

# Footer
st.markdown("---")
st.markdown("**📊 Chiefdom GPS Dashboard | School-Based Distribution Analysis**")
